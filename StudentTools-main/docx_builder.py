"""
docx_builder.py — Converts Markdown text to a professional .docx document.

Uses python-docx to create Word documents from AI-generated Markdown.
Parses Markdown headings, paragraphs, bold/italic, lists, tables, and blockquotes.
"""

import io
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def markdown_to_docx(markdown_text: str, topic: str = "") -> bytes:
    """
    Convert a Markdown string into a professionally formatted .docx file.
    Returns the .docx as bytes.
    """
    doc = Document()
    
    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # --- Default Font Style ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Configure heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        if level == 1:
            heading_style.font.size = Pt(24)
        elif level == 2:
            heading_style.font.size = Pt(16)
        elif level == 3:
            heading_style.font.size = Pt(13)
    
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # --- Headings ---
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        elif stripped.startswith('# '):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        
        # --- Table Detection ---
        if '|' in stripped and i + 1 < len(lines):
            # Check if this is a table (has header separator on next line)
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ''
            if re.match(r'^[\|\s\-:]+$', next_line):
                # Parse table
                table_lines = []
                j = i
                while j < len(lines) and '|' in lines[j].strip():
                    table_lines.append(lines[j].strip())
                    j += 1
                
                if len(table_lines) >= 3:  # header + separator + at least 1 row
                    _add_table(doc, table_lines)
                    i = j
                    continue
        
        # --- Blockquote ---
        if stripped.startswith('> '):
            quote_text = stripped[2:].strip()
            # Collect multi-line blockquotes
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith('> '):
                quote_text += ' ' + lines[j].strip()[2:].strip()
                j += 1
            
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            # Add a left border effect via indentation and italic
            run = p.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(11)
            i = j
            continue
        
        # --- Unordered List ---
        if re.match(r'^[\-\*]\s', stripped):
            bullet_text = re.sub(r'^[\-\*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, bullet_text)
            i += 1
            continue
        
        # --- Ordered List ---
        if re.match(r'^\d+\.\s', stripped):
            list_text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, list_text)
            i += 1
            continue
        
        # --- Horizontal Rule ---
        if re.match(r'^[\-\*_]{3,}$', stripped):
            # Add a subtle separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # --- Regular Paragraph ---
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_formatted_text(p, stripped)
        i += 1
    
    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_formatted_text(paragraph, text: str):
    """Parse inline Markdown formatting (bold, italic) and add runs to paragraph."""
    # Pattern matches **bold**, *italic*, and regular text
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    
    for match in re.finditer(pattern, text):
        bold_italic = match.group(2)
        bold = match.group(3)
        italic = match.group(4)
        normal = match.group(5)
        
        if bold_italic:
            run = paragraph.add_run(bold_italic)
            run.bold = True
            run.italic = True
        elif bold:
            run = paragraph.add_run(bold)
            run.bold = True
        elif italic:
            run = paragraph.add_run(italic)
            run.italic = True
        elif normal:
            paragraph.add_run(normal)


def _add_table(doc, table_lines: list):
    """Parse Markdown table lines and add a formatted table to the document."""
    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
    
    # Parse data rows (skip separator line at index 1)
    rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
    
    if not headers or not rows:
        return
    
    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < len(headers):  # Prevent index errors
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    
    # Add spacing after table
    doc.add_paragraph()
